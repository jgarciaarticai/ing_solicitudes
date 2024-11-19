from docx import Document
import re
import logging
from logging_config import setup_logging



class DocumentProcessor:
    def __init__(self, file_path):
        self.logger = logging.getLogger(__name__)
        self.file_path = file_path
        self.document = None
        self.sections = {}

    def load_document(self):
        """Carga el documento Word y lo prepara para el procesamiento."""
        try:
            self.document = Document(self.file_path)
            self.logger.info("Documento cargado con éxito.")
        except Exception as e:
            self.logger.error(f"Error al cargar el documento: {e}")

    def extract_index_titles(self, keywords):
        """Busca títulos en el índice que contengan palabras clave y extrae el número de página."""
        index_titles = []
        pattern = '|'.join(re.escape(keyword) for keyword in keywords)
        regex = re.compile(pattern, re.IGNORECASE)

        # Patrón para detectar títulos en el índice
        index_pattern = re.compile(
            r"^(?P<title>.+?)\s+(?P<page>\d+)$"  # Texto del título seguido por un número de página
        )

        self.logger.info("Analizando el índice para encontrar títulos relevantes...")

        for para in self.document.paragraphs:
            text = para.text.strip()

            match = index_pattern.match(text)
            if match and regex.search(match.group("title")):
                title = match.group("title")
                page = int(match.group("page"))

                # Limpieza del título: eliminar numeración inicial, tabulaciones y caracteres innecesarios
                clean_title = re.sub(r"^\d+(\.\d+)*\s*", "", title).strip()
                clean_title = re.sub(r"^\.\s*", "", clean_title)  # Eliminar punto y tabulación iniciales
                self.logger.info(f"Título válido del índice encontrado: {text} -> Título limpio: {clean_title}, Página: {page}")
                index_titles.append((clean_title, page))

        self.logger.info(f"Títulos identificados en el índice: {index_titles}")
        return index_titles

    def find_section_content(self, title):
        """Busca el contenido de una sección basándose en su título y estilo."""
        content = []
        capturing = False

        self.logger.info(f"Buscando título '{title}' en el documento...")

        for para in self.document.paragraphs:
            text = para.text.strip()
            style = para.style.name

            # Iniciar captura si encontramos el título con el estilo de título
            if title in text and style.startswith("ARTICA") and not capturing:
                capturing = True
                self.logger.info(f"Título encontrado: '{text}' con estilo {style}")
                continue

            # Detener la captura si encontramos otro título con estilo de título
            if capturing:
                if style.startswith("ARTICA") and title not in text:
                    self.logger.info(f"Fin de la sección para el título '{title}' detectado.")
                    break
                content.append(text)

        if capturing:
            self.logger.info(f"Contenido capturado para '{title}': {content[:5]}...")
        else:
            self.logger.error(f"No se pudo capturar contenido para '{title}'.")

        return "\n".join(content)

    def identify_sections(self, keywords):
        """Identifica las secciones basadas en el índice y recupera el contenido correspondiente."""
        if self.document is None:
            raise ValueError("El documento no ha sido cargado.")

        # Paso 1: Extraer títulos del índice
        index_titles = self.extract_index_titles(keywords)
        if not index_titles:
            self.logger.info("No se encontraron títulos en el índice con las palabras clave.")
            return

        self.logger.info("Buscando contenido correspondiente a los títulos detectados...")
        for title, page in index_titles:
            content = self.find_section_content(title)
            if content:
                self.sections[title] = content
                self.logger.info(f"\nSección encontrada: {title}")
                self.logger.info(f"Contenido completo:\n{content}\n")
            else:
                self.logger.info(f"No se encontró contenido para el título: {title}")

        self.logger.info(f"Secciones identificadas: {len(self.sections)}")

    def get_sections(self):
        """Devuelve las secciones identificadas para revisión o exportación."""
        return self.sections
