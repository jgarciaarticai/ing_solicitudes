from docx import Document
import os
import re
import logging
from docx.shared import Inches
from docx.oxml.ns import qn
from PIL import Image
from io import BytesIO
import base64
from docx.oxml import OxmlElement
from docx.shared import Pt

class DocumentExporter:
    def __init__(self, sections, output_dir):
        self.sections = sections
        self.output_dir = output_dir
        self.logger = logging.getLogger(__name__)  # Logger para el exportador


    def create_or_get_style(self, doc, style_name):
        """Crea o obtiene un estilo de párrafo llamado style_name."""
        styles = doc.styles

        if style_name in styles:  # Si el estilo ya existe, lo obtenemos
            return styles[style_name]

        # Crear nuevo estilo si no existe
        new_style = styles.add_style(style_name, 1)  # 1 para estilo de párrafo
        new_style.font.name = "Arial"  # Puedes cambiar esto a la fuente que desees
        new_style.font.size = Pt(14)   # Cambia el tamaño si lo necesitas
        new_style.font.bold = True     # Opcional

        # Configurar opciones adicionales, si es necesario
        return new_style


    def export_section(self, title, export_format="docx"):
        """Exporta una sección específica basada en el título."""
        try:
            if title not in self.sections:
                self.logger.error(f"Título no encontrado: {title}")
                return

            # Crear directorio si no existe
            os.makedirs(self.output_dir, exist_ok=True)

            # Crear documento
            new_doc = Document()
            custom_style = self.create_or_get_style(new_doc, style_name="ARTICA 6")
            heading = new_doc.add_paragraph(style=custom_style)
            heading_run = heading.add_run(title)

            for element in self.sections[title]:
                if element["type"] == "text":
                    para = new_doc.add_paragraph()
                    self._add_formatted_text(para, element["content"])
                elif element["type"] == "image":
                    image_run = element["content"]  # El objeto `Run` que contiene la imagen
                    para = new_doc.add_paragraph()  # Crear un párrafo para la imagen
                    self._add_image_to_document(para, image_run)

            # Guardar archivo
            sanitized_title = re.sub(r'[\\/*?:"<>|]', "", title)  # Evitar caracteres no válidos en nombres de archivo
            file_name = f"{self.output_dir}/{sanitized_title}.{export_format}"
            new_doc.save(file_name)
            self.logger.info(f"Documento exportado: {file_name}")
        except Exception as e:
            self.logger.exception(f"Error al exportar la sección '{title}': {e}")


    def _add_formatted_text(self, para, text_elements):
        """Agrega texto con formato al párrafo."""
        try:
            for element in text_elements:
                if isinstance(element, str):
                    # Si el elemento es un string plano, agrégalo directamente
                    para.add_run(element)
                elif isinstance(element, dict) and "text" in element:
                    # Si el elemento es un dict con formato, aplica los estilos
                    run = para.add_run(element["text"])
                    run.bold = element.get("bold", False)
                    run.italic = element.get("italic", False)
                    run.underline = element.get("underline", False)
                else:
                    self.logger.warning(f"Formato inesperado en texto: {element}")
        except Exception as e:
            self.logger.exception(f"Error al agregar texto con formato: {e}")


    def _add_image_to_document(self, paragraph, image_run):
        """Agrega una imagen extraída desde un objeto `Run` al párrafo."""
        try:
            # Extraer los datos binarios de la imagen desde el `Run`
            blip_elements = image_run._element.xpath(".//a:blip")
            if not blip_elements:
                self.logger.warning("No se encontraron datos de imagen en el `Run`.")
                return

            embed = blip_elements[0].get(qn("r:embed"))
            part = image_run.part.related_parts[embed]
            image_data = part.blob

            # Guardar temporalmente la imagen
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                temp_file.write(image_data)
                temp_file.flush()
                temp_image_path = temp_file.name

            # Agregar la imagen al párrafo en un nuevo `Run`
            run = paragraph.add_run()
            run.add_picture(temp_image_path, width=Inches(4))

            # Eliminar el archivo temporal
            os.remove(temp_image_path)

        except Exception as e:
            self.logger.exception(f"Error al agregar imagen al documento: {e}")


    def export_all_sections(self, export_format="docx"):
        """Exporta todas las secciones disponibles."""
        if not self.sections:
            self.logger.warning("No hay secciones para exportar.")
            return

        try:
            os.makedirs(self.output_dir, exist_ok=True)
            for title in self.sections:
                self.export_section(title, export_format=export_format)
            self.logger.info("Todas las secciones se han exportado correctamente.")
        except Exception as e:
            self.logger.exception(f"Error al exportar las secciones: {e}")
