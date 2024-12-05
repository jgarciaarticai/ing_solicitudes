from docx import Document
from pathlib import Path
import shutil
import logging
import pandas as pd
from procesar import DocumentProcessor
from exportar import DocumentExporter
from docx.shared import Inches
from docx.oxml import OxmlElement

class ContentInserter:
    def __init__(self, input_dir, config_dir, exporter: DocumentExporter):
        self.input_dir = Path(input_dir).resolve()
        self.config_dir = Path(config_dir).resolve()
        self.exporter = exporter  # Instancia de DocumentExporter
        self.logger = logging.getLogger(__name__)


    def process_files(self, mapping_file, proyecto_menor):
        """Procesa todos los archivos .docx en la carpeta de origen."""
        try:
            mapping = pd.read_excel(mapping_file)

            for doc_file in self.input_dir.glob("*.docx"):
                keyword = doc_file.stem.strip()

                # Filtrar el mapeo por keyword
                row = mapping[mapping['keyword'].str.strip() == keyword]
                if row.empty:
                    self.logger.warning(f"No se encontró un mapeo para '{keyword}'.")
                    continue

                row = row.iloc[0]
                presentacion = row['presentacion'].strip() if pd.notna(row.get('presentacion')) else None
                memoria = row['memoria'].strip() if pd.notna(row.get('memoria')) else None

                # Determinar la plantilla adecuada
                destino_template = presentacion if proyecto_menor else memoria
                if not destino_template:
                    self.logger.warning(f"No se especificó una plantilla para '{keyword}' (proyecto menor: {proyecto_menor}).")
                    continue

                config_folder = "presentaciones" if proyecto_menor else "memorias"
                template_path = (self.config_dir / config_folder / destino_template).resolve()
                self.logger.info(f"Validando plantilla en: {template_path}")

                if not template_path.exists():
                    self.logger.warning(f"Plantilla no encontrada: {template_path}")
                    continue

                destino_path = (self.input_dir / destino_template).resolve()
                shutil.copy(template_path, destino_path)
                self.logger.info(f"Plantilla copiada desde {template_path} a {destino_path}")

                # Procesar contenido del documento
                self._insert_content(doc_file, destino_path)

        except Exception as e:
            self.logger.exception(f"Error al procesar archivos: {e}")


    def _insert_content(self, doc_origen_path, doc_destino_path):
        """Inserta el contenido del documento de origen en el documento de destino."""
        try:
            doc_processor = DocumentProcessor(doc_origen_path)
            doc_processor.load_document()

            # Extraer el título desde el documento de origen (el nombre del archivo es el título)
            titulo_origen = Path(doc_origen_path).stem

            # Obtener el contenido de la sección específica usando find_section_content
            section_content = doc_processor.find_section_content(titulo_origen)

            if not section_content:
                self.logger.warning(f"No se encontró contenido para la sección '{titulo_origen}'.")
                return

            # Abrir el documento de destino
            doc_destino = Document(doc_destino_path)

            # Buscar el título correspondiente en el cuerpo del documento con estilo que empiece con "ARTICA"
            encontrado = False
            for paragraph in doc_destino.paragraphs:
                if titulo_origen.lower() in paragraph.text.lower() and paragraph.style.name.startswith("ARTICA"):
                    encontrado = True
                    self.logger.info(f"Apartado encontrado en el cuerpo del documento: {paragraph.text}")

                    # Insertar contenido justo después del título encontrado
                    self._insert_text_and_images(section_content, doc_destino, paragraph)
                    break

            if not encontrado:
                self.logger.warning(f"Apartado '{titulo_origen}' no encontrado en el cuerpo del documento destino.")
            else:
                # Guardar los cambios en el documento de destino
                doc_destino.save(doc_destino_path)
                self.logger.info(f"Contenido de '{doc_origen_path}' insertado en '{doc_destino_path}'.")

        except Exception as e:
            self.logger.exception(f"Error al insertar contenido de '{doc_origen_path}' en '{doc_destino_path}': {e}")


    def _insert_text_and_images(self, content, doc_destino, paragraph):
        """Inserta texto e imágenes después de un párrafo específico."""
        try:
            # Obtener el elemento XML del párrafo actual
            current_element = paragraph._element

            for element in content:
                if element["type"] == "text":
                    # Crear un nuevo párrafo después del párrafo actual
                    new_para = doc_destino.add_paragraph()
                    self.exporter._add_formatted_text(new_para, element["content"])

                    # Mover el nuevo párrafo después del párrafo actual en el XML
                    current_element.addnext(new_para._element)
                    current_element = new_para._element  # Actualizar la referencia

                elif element["type"] == "image":
                    # Crear un párrafo para la imagen
                    new_para = doc_destino.add_paragraph()
                    self.exporter._add_image_to_document(new_para, element["content"])

                    # Mover el párrafo de la imagen después del párrafo actual
                    current_element.addnext(new_para._element)
                    current_element = new_para._element  # Actualizar la referencia

            self.logger.info("Texto e imágenes insertados correctamente.")

        except Exception as e:
            self.logger.exception(f"Error al insertar texto e imágenes: {e}")
