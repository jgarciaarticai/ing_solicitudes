from docx import Document
import re
import logging
from logging_config import setup_logging

class DocumentProcessor:
    def __init__(self, file_path):
        self.logger = logging.getLogger(__name__)
        self.file_path = file_path
        self.document = None
        self.sections = {}


    def load_document(self):
        """Carga el documento Word y lo prepara para el procesamiento."""
        try:
            self.document = Document(self.file_path)
            self.logger.info("Documento cargado con éxito.")
        except Exception as e:
            self.logger.exception(f"Error al cargar el documento: {e}")
            raise


    def extract_index_titles(self, keywords):
        """Busca títulos en el índice que contengan palabras clave y extrae el número de página."""
        try:
            index_titles = []
            pattern = '|'.join(re.escape(keyword) for keyword in keywords)
            regex = re.compile(pattern, re.IGNORECASE)

            # Patrón para detectar títulos en el índice
            index_pattern = re.compile(
                r"^(?P<title>.+?)\s+(?P<page>\d+)$"  # Texto del título seguido por un número de página
            )

            self.logger.info("Analizando el índice para encontrar títulos relevantes...")

            for para in self.document.paragraphs:
                text = para.text.strip()

                match = index_pattern.match(text)
                if match and regex.search(match.group("title")):
                    title = match.group("title")
                    page = int(match.group("page"))

                    # Limpieza del título
                    clean_title = re.sub(r"^\d+(\.\d+)*\s*", "", title).strip()
                    clean_title = re.sub(r"^\.\s*", "", clean_title)  # Eliminar punto y tabulación iniciales
                    self.logger.info(f"Título válido del índice encontrado: {text} -> Título limpio: {clean_title}, Página: {page}")
                    index_titles.append((clean_title, page))

            self.logger.info(f"Títulos identificados en el índice: {index_titles}")
            return index_titles
        except Exception as e:
            self.logger.exception(f"Error al analizar el índice: {e}")
            return []


    def find_section_content(self, title):
        """Busca el contenido de una sección basándose en su título y estilo."""
        try:
            content = []
            capturing = False
            self.logger.info(f"Buscando título '{title}' en el documento...")

            for para in self.document.paragraphs:
                text = para.text.strip()
                style = para.style.name

                # Iniciar captura si encontramos el título con el estilo adecuado
                if title in text and style.startswith("ARTICA") and not capturing:
                    capturing = True
                    self.logger.info(f"Título encontrado: '{text}' con estilo {style}")
                    content.append({"type": "text", "content": [{"text": text}]})
                    continue

                # Detener la captura si encontramos otro título con estilo de título
                if capturing:
                    if style.startswith("ARTICA") and title not in text:
                        self.logger.info(f"Fin de la sección para el título '{title}' detectado.")
                        break

                    # Capturar texto con formato
                    paragraph_content = []
                    for run in para.runs:
                        paragraph_content.append({
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                        })

                    if paragraph_content:
                        content.append({"type": "text", "content": paragraph_content})

                    # Detectar imágenes en el párrafo
                    for run in para.runs:
                        if run._element.xpath(".//w:drawing"):
                            self.logger.info("Imagen encontrada.")
                            content.append({"type": "image", "content": run})  # Guardar el `Run` de la imagen


            if capturing:
                self.logger.info(f"Contenido capturado para '{title}': {content[:2]}...")
            else:
                self.logger.warning(f"No se pudo capturar contenido para '{title}'.")

            return content
        except Exception as e:
            self.logger.exception(f"Error al buscar contenido para '{title}': {e}")
            return []


    def identify_sections(self, keywords):
        """Identifica las secciones basadas en el índice y recupera el contenido correspondiente."""
        try:
            if self.document is None:
                raise ValueError("El documento no ha sido cargado.")

            # Extraer títulos del índice
            index_titles = self.extract_index_titles(keywords)
            if not index_titles:
                self.logger.warning("No se encontraron títulos en el índice con las palabras clave.")
                return

            self.logger.info("Buscando contenido correspondiente a los títulos detectados...")
            for title, page in index_titles:
                content = self.find_section_content(title)
                if content:
                    self.sections[title] = content
                    self.logger.info(f"Sección encontrada: {title}")
                else:
                    self.logger.warning(f"No se encontró contenido para el título: {title}")

            self.logger.info(f"Secciones identificadas: {len(self.sections)}")
        except Exception as e:
            self.logger.exception(f"Error al identificar secciones: {e}")


    def get_sections(self):
        """Devuelve las secciones identificadas para revisión o exportación."""
        try:
            return self.sections
        except Exception as e:
            self.logger.exception(f"Error al obtener las secciones: {e}")
            return {}
